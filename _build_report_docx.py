# -*- coding: utf-8 -*-
"""Генерация отчёта ПР9-12 под проект «бронирование переговорок»."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_centered(doc: Document, text: str, bold: bool = False, size: int | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def main() -> None:
    root = Path(__file__).resolve().parent
    out = root / "ИПОРиПИС_СемейкинСВ_ПР9-12_отчет.docx"

    doc = Document()

    for line in [
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "Федеральное государственное бюджетное образовательное учреждение",
        "высшего образования",
        "«МИРЭА – Российский технологический университет»",
        "РТУ МИРЭА",
        "Институт информационных технологий",
        "Кафедра инструментального и прикладного программного обеспечения",
        "(ИППО)",
    ]:
        add_centered(doc, line)

    doc.add_paragraph()
    add_centered(doc, "ОТЧЁТ", bold=True, size=14)
    add_centered(doc, "ПО ПРАКТИЧЕСКИМ РАБОТАМ №9–12", bold=True, size=14)
    doc.add_paragraph()
    add_centered(
        doc,
        "по дисциплине «Инструментальное программное обеспечение разработки и "
        "проектирования информационных систем»",
    )
    doc.add_paragraph()
    add_centered(doc, "Выполнил студент группы _________________ Семейкин С. В.")
    add_centered(doc, "Принял преподаватель _________________________________")
    doc.add_paragraph()
    add_centered(doc, "Практическая работа выполнена «___» _____________ 2026 г.")
    add_centered(doc, "«Зачтено» «___» _____________ 2026 г.")
    doc.add_paragraph()
    add_centered(doc, "Москва 2026")

    doc.add_page_break()
    add_heading(doc, "Содержание", 1)
    for t in [
        "Цель цикла и выбранная тема",
        "Практическая работа №9",
        "Практическая работа №10",
        "Практическая работа №11",
        "Практическая работа №12",
        "Заключение",
    ]:
        doc.add_paragraph(t, style="List Number")

    doc.add_page_break()
    add_heading(doc, "Цель цикла и выбранная тема", 1)
    add_body(
        doc,
        "Цель цикла – освоить на практике современные методы и инструменты создания "
        "программных продуктов: от архитектурного проектирования в нотации C4 до "
        "развёртывания в Kubernetes и подключения мониторинга (Prometheus, Grafana), "
        "с использованием ИИ-ассистентов на отдельных этапах.",
    )
    add_body(
        doc,
        "Индивидуальная тема – веб-приложение для бронирования переговорных комнат в "
        "организации: поиск свободных слотов по времени и вместимости, создание броней "
        "с проверкой пересечений, приглашение зарегистрированных пользователей на встречу "
        "(поиск по email) с отправкой приглашения на почту при настроенном SMTP, роли "
        "«сотрудник» и «офис-администратор», хранение данных в PostgreSQL, кэш и блокировки "
        "на Redis, вход по email и паролю с выдачей JWT, единая точка входа через API Gateway.",
    )

    # --- ПР9 ---
    doc.add_page_break()
    add_heading(doc, "Практическая работа №9", 1)
    add_heading(doc, "Тема и цель работы", 2)
    add_body(
        doc,
        "Тема – веб-приложение для бронирования переговорок. Цель работы – освоить "
        "методологию C4 (уровни контекст, контейнеры, компоненты), сформулировать "
        "Problem Statement, подготовить диаграммы в PlantUML с библиотекой C4-PlantUML "
        "и выполнить критический разбор результатов генерации с помощью ИИ.",
    )
    add_heading(doc, "Постановка задачи (Problem Statement)", 2)
    add_body(
        doc,
        "Компании требуется веб-приложение: сотрудники ищут свободные переговорные "
        "комнаты и оформляют бронирования; организатор встречи может искать других "
        "пользователей системы по email, добавлять их к брони и направлять приглашение "
        "на почту (через внешний SMTP). Офис-администратор управляет справочником "
        "комнат. Система не должна допускать пересечения броней по одной комнате; "
        "необходима история для анализа загрузки. В перспективе возможна интеграция "
        "с корпоративным SSO; в рамках MVP реализованы регистрация, вход и JWT внутри приложения.",
    )
    add_heading(doc, "Используемые средства", 2)
    add_body(
        doc,
        "Нотация C4, синтаксис C4-PlantUML, исходники диаграмм в каталоге "
        "practice1/diagrams (*.puml). При подготовке черновиков использовались "
        "подсказки ИИ; итоговые границы системы, набор контейнеров и подписи связей "
        "согласованы с фактическим стеком репозитория (React/Vite, FastAPI, PostgreSQL, Redis, "
        "опциональная отправка почты smtplib при заданных SMTP-переменных).",
    )
    add_heading(doc, "Контекстная диаграмма (C1)", 2)
    add_body(
        doc,
        "На уровне контекста отражены сотрудник и офис-администратор как пользователи, "
        "разрабатываемая система как единый продукт и внешний канал доставки почты "
        "(SMTP relay) для приглашений участников встречи. Детализация внутренних сервисов "
        "вынесена на уровень контейнеров (C2). "
        "Контекстная диаграмма представлена на Рисунке 1 (файл practice1/diagrams/context.puml).",
    )
    add_heading(doc, "Диаграмма контейнеров (C2)", 2)
    add_body(
        doc,
        "Показаны развёртываемые части: Web UI (React, Vite, nginx), API Gateway (FastAPI), "
        "Booking Service (FastAPI, SQLAlchemy), PostgreSQL, Redis; потоки HTTPS/HTTP/SQL/RESP "
        "и исходящий SMTP из booking-service к почтовому relay для приглашений. "
        "Диаграмма представлена на Рисунке 2 (practice1/diagrams/container.puml).",
    )
    add_heading(doc, "Диаграмма компонентов (C3)", 2)
    add_body(
        doc,
        "Раскрыт контейнер Booking Service: REST Router (в т.ч. поиск пользователей и "
        "управление участниками брони), слой безопасности (JWT, bcrypt), бизнес-логика, "
        "компонент отправки приглашений по SMTP (smtplib), репозиторий, адаптер Redis, "
        "экспорт метрик Prometheus. "
        "Диаграмма представлена на Рисунке 3 (practice1/diagrams/component.puml).",
    )
    add_heading(doc, "Критический анализ и доработка диаграмм", 2)
    add_body(
        doc,
        "Типичные расхождения при генерации ИИ: лишние прямые связи пользователя с БД, "
        "несогласованные названия между уровнями, избыточные микросервисы. Вручную "
        "уточнены границы системы, добавлен Redis как внешнее хранилище кэша, "
        "разделены ответственности Router/Service/Repository и слой JWT, отражены "
        "поток приглашения участников и связь с внешним SMTP. "
        "Таблица критического анализа приведена в practice1/README.md.",
    )
    add_heading(doc, "Вывод", 2)
    add_body(
        doc,
        "Получены три согласованных уровня C1–C3, оформлен Problem Statement и "
        "электронные артефакты. ИИ целесообразно использовать для черновых PlantUML; "
        "финальная архитектура проверена вручную на соответствие стеку MVP.",
    )

    # --- ПР10 ---
    doc.add_page_break()
    add_heading(doc, "Практическая работа №10", 1)
    add_heading(doc, "Тема и цель работы", 2)
    add_body(
        doc,
        "Тема совпадает с практикой №9. Цель – реализовать MVP из не менее двух "
        "микросервисов, поднять стек в Docker Compose, обеспечить PostgreSQL и Redis, "
        "межсервисное взаимодействие через HTTP, автотесты и описание использования ИИ.",
    )
    add_heading(doc, "Архитектура MVP", 2)
    add_body(
        doc,
        "Декомпозиция согласована с диаграммой C2 из практики №9. "
        "Микросервис 1 – API Gateway (practice2/services/gateway): FastAPI, проксирование "
        "REST-запросов к booking-service (в т.ч. /api/users/search и участники брони), "
        "передача заголовка Authorization, эндпоинт /metrics. "
        "Микросервис 2 – Booking Service (practice2/services/booking_service): бизнес-логика "
        "комнат и бронирований, связь «бронь – участники», JWT и роли пользователя, "
        "SQLAlchemy + PostgreSQL, Redis для кэша и блокировок, /metrics, отправка "
        "приглашений через smtplib при заданных переменных SMTP. Frontend – SPA на React "
        "(practice2/frontend), сценарий «Участники встречи» для организатора, обращение к gateway.",
    )
    add_heading(doc, "Модель данных и безопасность", 2)
    add_body(
        doc,
        "В PostgreSQL хранятся пользователи (в т.ч. bootstrap-администратор), переговорные "
        "комнаты, бронирования и таблица участников брони (связь many-to-many бронь–пользователь). "
        "Секреты (JWT_SECRET_KEY, пароли БД, пароль bootstrap-админа, при необходимости SMTP) "
        "задаются через practice2/.env (шаблон – practice2/.env.example), в репозиторий не "
        "коммитится; docker-compose.yml подставляет переменные окружения.",
    )
    add_heading(doc, "Инфраструктура", 2)
    add_body(
        doc,
        "Файл practice2/docker-compose.yml описывает сервисы frontend, gateway, "
        "booking-service, postgres, redis. Перед первым запуском необходимо скопировать "
        ".env.example в .env и при необходимости изменить значения (включая опциональные SMTP_* "
        "для почтовых приглашений). Проверка: "
        "docker compose up -d --build; UI на порту 3000, gateway на 8000.",
    )
    add_heading(doc, "Межсервисное взаимодействие", 2)
    add_body(
        doc,
        "Браузер обращается к gateway по HTTPS (в локальной среде – через опубликованный "
        "порт). Gateway пересылает запросы на booking-service по внутренней сети Compose; "
        "токен JWT проверяется на стороне booking-service. Поиск пользователей и добавление "
        "участников к брони выполняются тем же контуром API; при успешном добавлении "
        "booking-service инициирует исходящее SMTP-соединение к relay, если задан SMTP_HOST.",
    )
    add_heading(doc, "Автотесты", 2)
    add_body(
        doc,
        "В каталоге practice2/tests размещены pytest-сценарии для booking-service и gateway "
        "(в т.ч. с подменой httpx для upstream и подменой отправки почты в тестах участников). "
        "Переменные окружения для тестов задаются "
        "в коде тестов до импорта приложения.",
    )
    add_heading(doc, "Использование ИИ при выполнении работы", 2)
    add_body(
        doc,
        "Использовалась среда Cursor: ускорены каркас сервисов, Dockerfile, compose, "
        "черновики тестов и README. Ручная доработка потребовалась для согласования "
        "DATABASE_URL и хостов внутри контейнеров, политики авторизации (админ/пользователь), "
        "выноса секретов в .env, сценария приглашений и обработки ошибок SMTP.",
    )
    add_heading(doc, "Вывод", 2)
    add_body(
        doc,
        "Реализован MVP бронирования переговорок на двух микросервисах с общей БД и Redis, "
        "поиском пользователей и участниками встречи, опциональной отправкой приглашений по SMTP, "
        "оркестрация через Docker Compose, автотесты. ИИ сократил время на шаблонный код; "
        "безопасность, целостность данных и корректность почтового сценария проверены вручную.",
    )

    # --- ПР11 ---
    doc.add_page_break()
    add_heading(doc, "Практическая работа №11", 1)
    add_heading(doc, "Тема и цель работы", 2)
    add_body(
        doc,
        "Цель – перенести MVP с Docker Compose в Kubernetes (Minikube): Deployment и "
        "Service для компонентов, Ingress для внешнего доступа, ConfigMap, Secret; "
        "освоить цикл kubectl apply и проверки состояния кластера.",
    )
    add_heading(doc, "Исходная система", 2)
    add_body(
        doc,
        "В практике №10 реализованы gateway и booking-service с поиском пользователей, "
        "участниками брони и опциональной отправкой приглашений по SMTP, а также PostgreSQL и Redis. "
        "В практике №11 изменён способ упаковки: манифесты в practice3/k8s/.",
    )
    add_heading(doc, "Микросервисы и образы", 2)
    add_body(
        doc,
        "В кластер выводятся образы gateway:latest и booking-service:latest (сборка из "
        "practice2, загрузка в Minikube), postgres:16-alpine, redis:7-alpine. "
        "Секреты оформляются отдельным файлом booking-secrets.yaml по шаблону "
        "booking-secrets.yaml.example (рабочая копия не коммитится); в Secret перенесены "
        "DATABASE_URL, JWT_SECRET_KEY, пароли PostgreSQL и bootstrap-администратора.",
    )
    add_heading(doc, "Структура манифестов", 2)
    add_body(
        doc,
        "Единый файл practice3/k8s/all-in-one.yaml содержит ConfigMap, деплойменты, "
        "сервисы и Ingress (хост myapp.local). Перед применением выполняется "
        "kubectl apply -f booking-secrets.yaml. Деплой booking-service получает "
        "JWT и строку подключения из Secret. Для работы почтовых приглашений в кластере "
        "в манифест booking-service дополнительно задаются (или расширяются вручную) "
        "переменные SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASSWORD, SMTP_FROM по аналогии с Docker Compose.",
    )
    add_heading(doc, "Развёртывание в Minikube", 2)
    add_body(
        doc,
        "Запуск Minikube, включение ingress, сборка и minikube image load для прикладных "
        "образов, применение Secret и all-in-one.yaml, ожидание готовности deployment "
        "и проверка kubectl get pods,svc,ingress. Подробная последовательность приведена в "
        "корневом README.md и practice3/PRACTICE3.md.",
    )
    add_heading(doc, "Проверка", 2)
    add_body(
        doc,
        "Контроль: состояние подов Running, доступ к API через Ingress или port-forward, "
        "проверка health, списка комнат, сценария поиска пользователей и добавления участника к брони. "
        "Рисунки с скриншотами kubectl и запросов "
        "к API включаются в печатную версию отчёта по месту (Рисунки 10–12 – по аналогии "
        "с методическими требованиями к оформлению).",
    )
    add_heading(doc, "Использование ИИ", 2)
    add_body(
        doc,
        "ИИ ускорил подготовку YAML и правок README; порядок применения манифестов, "
        "имена сервисов в DATABASE_URL и согласование Secret проверялись вручную на кластере.",
    )
    add_heading(doc, "Вывод", 2)
    add_body(
        doc,
        "Приложение развёрнуто в Minikube с использованием Deployment, Service, Ingress, "
        "ConfigMap и вынесенного в отдельный файл Secret; обеспечен внешний доступ к gateway.",
    )

    # --- ПР12 ---
    doc.add_page_break()
    add_heading(doc, "Практическая работа №12", 1)
    add_heading(doc, "Тема и цель работы", 2)
    add_body(
        doc,
        "Цель – подключить наблюдаемость: метрики в формате Prometheus на /metrics, "
        "сбор через ServiceMonitor, дашборд Grafana и нагрузочный прогон.",
    )
    add_heading(doc, "Исходная система", 2)
    add_body(
        doc,
        "Опора на кластер из практики №11. Доработан код gateway и booking-service: "
        "счётчики и гистограммы HTTP по путям (включая новые маршруты поиска и участников), "
        "бизнес-счётчик bookings_created_total в booking-service.",
    )
    add_heading(doc, "Стек мониторинга", 2)
    add_body(
        doc,
        "Prometheus и Grafana устанавливаются Helm-чартом kube-prometheus-stack. "
        "В practice4/monitoring размещены servicemonitors.yaml и grafana-dashboard.json; "
        "нагрузочный сценарий – practice4/load_test.py.",
    )
    add_heading(doc, "Инструментирование", 2)
    add_body(
        doc,
        "Gateway экспортирует gateway_http_requests_total и "
        "gateway_http_request_duration_seconds; booking-service – booking_http_requests_total, "
        "booking_http_request_duration_seconds и bookings_created_total. Эндпоинт /metrics "
        "доступен для опроса Prometheus.",
    )
    add_heading(doc, "Сбор метрик и нагрузка", 2)
    add_body(
        doc,
        "После kubectl apply для ServiceMonitor цели приложения отображаются в Prometheus "
        "как UP. В Grafana импортируется дашборд с панелями RPS, p95 latency и скорости "
        "создания бронирований. Нагрузка: python practice4/load_test.py (порядка 150 запросов). "
        "Рисунки 13–15 – скриншоты Prometheus, Grafana до и после нагрузки.",
    )
    add_heading(doc, "Использование ИИ", 2)
    add_body(
        doc,
        "ИИ использовался для черновиков метрик, YAML и описания в PRACTICE4.md; "
        "проверка scrape, проброс портов к Grafana и адекватность графиков – вручную.",
    )
    add_heading(doc, "Вывод", 2)
    add_body(
        doc,
        "Микросервисы инструментированы, метрики собираются Prometheus, визуализируются в "
        "Grafana; нагрузочный тест демонстрирует реакцию показателей.",
    )

    # Заключение
    doc.add_page_break()
    add_heading(doc, "Заключение", 1)
    add_body(
        doc,
        "Прохождение цикла практик №9–12 с опорой на ИИ-ассистента позволило быстрее "
        "получить черновики диаграмм, кода и манифестов, но финальная согласованность "
        "архитектуры, секретов, сетевых имён и поведения под нагрузкой остаётся зоной "
        "ответственности разработчика.",
    )
    add_body(
        doc,
        "ИИ наиболее полезен на этапах с выраженным шаблоном: PlantUML, docker-compose, "
        "типовые Kubernetes-ресурсы, заготовки метрик. Менее допустим без контроля "
        "контекст хранения секретов и границы доверия между сервисами.",
    )
    add_body(
        doc,
        "Дополнительные элементы, приближающие проект к промышленной практике: "
        "разделение gateway и доменной логики, Redis для конкурентного доступа к слотам, "
        "вынесение секретов в .env и Kubernetes Secret, приглашения участников встречи по SMTP, "
        "экспорт метрик и дашборд в Grafana.",
    )

    doc.save(out)
    print("Wrote", out)


if __name__ == "__main__":
    main()
